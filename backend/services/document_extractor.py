"""
Document Extraction Service - supports PDF and Word (.docx) files.

Note: PDF text extraction is now handled by Gemini API (gemini_extractor.py).
This service is kept for backward compatibility and for Word documents only.
"""

import io
import logging
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, Any, List, Tuple

logger = logging.getLogger(__name__)


def _extract_docx_text(file_bytes: bytes) -> List[str]:
    """Extract text paragraphs from a .docx file using only stdlib (zipfile + xml).
    No python-docx or lxml needed. Returns list of paragraph strings."""
    paragraphs = []
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            # Main document body
            if "word/document.xml" in zf.namelist():
                with zf.open("word/document.xml") as f:
                    tree = ET.parse(f)
                    root = tree.getroot()
                    # Word XML namespace
                    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                    for para in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
                        texts = []
                        for run in para.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                            if run.text:
                                texts.append(run.text)
                        line = "".join(texts).strip()
                        if line:
                            paragraphs.append(line)
    except (zipfile.BadZipFile, ET.ParseError, KeyError) as e:
        logger.warning(f"Failed to extract docx text with stdlib: {e}")
        # Fallback to python-docx if available
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
        except ImportError:
            raise ValueError("Could not extract text from Word document")
    return paragraphs


class DocumentExtractor:
    """Extracts text from PDF or Word documents.
    
    Note: For PDFs, use GeminiExtractor instead for better accuracy.
    This service is primarily for Word documents.
    """

    def extract(self, file_bytes: bytes, file_name: str) -> Dict[str, Any]:
        """
        Auto-detect file type and extract text with metadata.

        Returns dict with:
          - full_text: str
          - blocks: list (with bbox for PDFs, paragraph-level for Word)
          - pages: list of page metadata
          - page_count: int
          - file_type: "pdf" | "docx"
          
        Note: For PDF files, use GeminiExtractor.analyze_pdf() instead.
        """
        lower_name = file_name.lower()

        if lower_name.endswith(".pdf"):
            logger.warning("DocumentExtractor.extract() called for PDF. Use GeminiExtractor.analyze_pdf() instead.")
            return self._extract_pdf(file_bytes)
        elif lower_name.endswith((".docx", ".doc")):
            return self._extract_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {file_name}")

    def _extract_pdf(self, file_bytes: bytes) -> Dict[str, Any]:
        """
        Extract from PDF using PyMuPDF with coordinate tracking (legacy).
        
        Deprecated: Use GeminiExtractor.analyze_pdf() for PDF analysis.
        This method is kept for fallback purposes only.
        """
        try:
            from services.pdf_extractor import PDFExtractor
            
            extractor = PDFExtractor()
            result = extractor.extract_text_with_coords(file_bytes)
            result["file_type"] = "pdf"
            logger.warning("Using legacy PDF extraction. Consider migrating to GeminiExtractor.")
            return result
        except Exception as e:
            logger.error(f"Legacy PDF extraction failed: {e}")
            raise

    def _extract_docx(self, file_bytes: bytes) -> Dict[str, Any]:
        """Extract from Word document using lightweight stdlib-based extraction."""
        paragraphs = _extract_docx_text(file_bytes)

        full_text_parts: List[str] = []
        blocks: List[Dict[str, Any]] = []
        paragraph_index = 0

        for text in paragraphs:
            full_text_parts.append(text)
            blocks.append({
                "text": text,
                "page": 0,
                "bbox": {
                    "x0": 72.0,
                    "y0": float(72 + paragraph_index * 20),
                    "x1": 540.0,
                    "y1": float(72 + paragraph_index * 20 + 16),
                },
            })
            paragraph_index += 1

        full_text = "\n".join(full_text_parts)

        return {
            "full_text": full_text,
            "blocks": blocks,
            "pages": [{"width": 612, "height": 792, "page_num": 0}],
            "page_count": 1,
            "file_type": "docx",
        }
