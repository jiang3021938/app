import logging
import base64
import json
import ast
import httpx as httpx_client
from datetime import datetime
from urllib.parse import urlencode, quote
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import Response
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
from sqlalchemy.ext.asyncio import AsyncSession

from core.database import get_db
from dependencies.auth import get_current_user
from schemas.auth import UserResponse
from services.lease_extraction import LeaseExtractionService, generate_ics_content
from services.pdf_extractor import PDFExtractor
from services.document_extractor import DocumentExtractor
from services.gemini_extractor import GeminiExtractor
from services.amendment_memo import AmendmentMemoService
from services.lease_comparison import LeaseComparisonService
from services.rent_benchmark import RentBenchmarkService
from services.executive_summary import ExecutiveSummaryService
from services.pdf_report import PDFReportGenerator
from services.documents import DocumentsService
from services.extractions import ExtractionsService
from services.user_credits import User_creditsService
from services.compliance_checker import ComplianceChecker
from services.storage import StorageService
from schemas.storage import FileUpDownRequest

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1/lease", tags=["lease"])


def safe_parse_json_or_literal(data: Any) -> Any:
    """Safely parse a string that could be JSON or Python literal.
    
    This replaces unsafe eval() calls with safe parsing methods.
    """
    if data is None:
        return None
    if not isinstance(data, str):
        return data
    
    # Try JSON first (most common case)
    try:
        return json.loads(data)
    except (json.JSONDecodeError, TypeError):
        pass
    
    # Try ast.literal_eval for Python literals (safe alternative to eval)
    try:
        return ast.literal_eval(data)
    except (ValueError, SyntaxError):
        pass
    
    # Return as-is if parsing fails
    return data


class AnalyzeRequest(BaseModel):
    document_id: int


class BatchAnalyzeRequest(BaseModel):
    document_ids: List[int]


class CalendarRequest(BaseModel):
    extraction_id: int


class AnalysisResponse(BaseModel):
    success: bool
    extraction_id: Optional[int] = None
    data: Optional[dict] = None
    compliance: Optional[dict] = None
    error: Optional[str] = None


class BatchAnalysisResponse(BaseModel):
    success: bool
    results: List[dict] = []
    total: int = 0
    completed: int = 0
    failed: int = 0


@router.post("/analyze", response_model=AnalysisResponse)
async def analyze_document(
    request: AnalyzeRequest,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Analyze a previously uploaded lease document"""
    try:
        # Get document
        doc_service = DocumentsService(db)
        document = await doc_service.get_by_id(request.document_id, current_user.id)
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Check user credits (admins bypass credit checks)
        is_admin = getattr(current_user, 'role', 'user') == 'admin'
        credits_service = User_creditsService(db)
        user_credits = await credits_service.get_by_field("user_id", current_user.id)
        
        if not user_credits:
            # Create initial credits for new user
            user_credits = await credits_service.create({
                "user_id": current_user.id,
                "free_credits": 1,
                "paid_credits": 0,
                "subscription_type": "none",
                "created_at": datetime.now(),
                "updated_at": datetime.now()
            }, current_user.id)
        
        # Check if user has credits
        total_credits = (user_credits.free_credits or 0) + (user_credits.paid_credits or 0)
        
        # Check subscription
        has_subscription = False
        if user_credits.subscription_type == "monthly" and user_credits.subscription_expires_at:
            if user_credits.subscription_expires_at > datetime.now():
                has_subscription = True
        
        if not is_admin and total_credits <= 0 and not has_subscription:
            raise HTTPException(
                status_code=402, 
                detail="No credits remaining. Please purchase more credits or subscribe."
            )
        
        # Update document status
        await doc_service.update(request.document_id, {"status": "processing"}, current_user.id)
        
        # Download file from storage
        storage = StorageService()
        download = await storage.create_download_url(
            FileUpDownRequest(bucket_name="lease-documents", object_key=document.file_key)
        )
        async with httpx_client.AsyncClient(timeout=120.0) as http_client:
            file_response = await http_client.get(download.download_url)
            file_bytes = file_response.content

        # Analyze PDF using Gemini API
        try:
            gemini_extractor = GeminiExtractor()
            analysis_result = await gemini_extractor.analyze_pdf(file_bytes, document.file_name)
            extracted_data = analysis_result["extracted_data"]
            full_text = analysis_result["full_text"]
            source_blocks = analysis_result["source_blocks"]  # Empty for Gemini
            pages_meta = analysis_result["pages_meta"]  # Empty for Gemini
            logger.info(f"Successfully analyzed PDF using Gemini API")
        except Exception as e:
            logger.error(f"Gemini analysis failed: {e}")
            await doc_service.update(request.document_id, {"status": "failed"}, current_user.id)
            return AnalysisResponse(
                success=False,
                error=f"PDF analysis failed: {str(e)}"
            )
        
        # Source mapping is not available with Gemini (no coordinate info)
        source_map = {}
        
        # Perform compliance check
        compliance_checker = ComplianceChecker(db)
        compliance_result = await compliance_checker.check_compliance(extracted_data)
        
        # Save extraction results
        extractions_service = ExtractionsService(db)
        extraction = await extractions_service.create({
            "user_id": current_user.id,
            "document_id": request.document_id,
            "tenant_name": extracted_data.get("tenant_name"),
            "landlord_name": extracted_data.get("landlord_name"),
            "property_address": extracted_data.get("property_address"),
            "monthly_rent": extracted_data.get("monthly_rent"),
            "security_deposit": extracted_data.get("security_deposit"),
            "lease_start_date": extracted_data.get("lease_start_date"),
            "lease_end_date": extracted_data.get("lease_end_date"),
            "renewal_notice_days": extracted_data.get("renewal_notice_days"),
            "pet_policy": extracted_data.get("pet_policy"),
            "late_fee_terms": extracted_data.get("late_fee_terms"),
            "risk_flags": json.dumps(extracted_data.get("risk_flags", [])),
            "compliance_data": json.dumps(compliance_result),
            "raw_extraction": json.dumps(extracted_data),
            "source_map": json.dumps(source_map) if source_map else None,
            "pages_meta": json.dumps(pages_meta) if pages_meta else None,
            "created_at": datetime.now()
        }, current_user.id)
        
        # Deduct credits (admins don't consume credits)
        if not is_admin:
            if user_credits.free_credits and user_credits.free_credits > 0:
                await credits_service.update(user_credits.id, {
                    "free_credits": user_credits.free_credits - 1,
                    "updated_at": datetime.now()
                }, current_user.id)
            elif user_credits.paid_credits and user_credits.paid_credits > 0:
                await credits_service.update(user_credits.id, {
                    "paid_credits": user_credits.paid_credits - 1,
                    "updated_at": datetime.now()
                }, current_user.id)
        
        # Update document status
        await doc_service.update(request.document_id, {"status": "completed"}, current_user.id)
        
        return AnalysisResponse(
            success=True,
            extraction_id=extraction.id,
            data=extracted_data,
            compliance=compliance_result
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Analysis error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze-batch", response_model=BatchAnalysisResponse)
async def analyze_documents_batch(
    request: BatchAnalyzeRequest,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Analyze multiple lease documents in batch"""
    try:
        results = []
        completed = 0
        failed = 0
        
        # Check user credits first (admins bypass)
        is_admin = getattr(current_user, 'role', 'user') == 'admin'
        credits_service = User_creditsService(db)
        user_credits = await credits_service.get_by_field("user_id", current_user.id)
        
        if not user_credits:
            user_credits = await credits_service.create({
                "user_id": current_user.id,
                "free_credits": 1,
                "paid_credits": 0,
                "subscription_type": "none",
                "created_at": datetime.now(),
                "updated_at": datetime.now()
            }, current_user.id)
        
        total_credits = (user_credits.free_credits or 0) + (user_credits.paid_credits or 0)
        has_subscription = False
        if user_credits.subscription_type == "monthly" and user_credits.subscription_expires_at:
            if user_credits.subscription_expires_at > datetime.now():
                has_subscription = True
        
        documents_to_process = len(request.document_ids)
        
        if not is_admin and total_credits < documents_to_process and not has_subscription:
            raise HTTPException(
                status_code=402,
                detail=f"Insufficient credits. You have {total_credits} credits but need {documents_to_process}."
            )
        
        doc_service = DocumentsService(db)
        extractions_service = ExtractionsService(db)
        extraction_service = LeaseExtractionService()
        compliance_checker = ComplianceChecker(db)
        
        for doc_id in request.document_ids:
            try:
                document = await doc_service.get_by_id(doc_id, current_user.id)
                
                if not document:
                    results.append({
                        "document_id": doc_id,
                        "success": False,
                        "error": "Document not found"
                    })
                    failed += 1
                    continue
                
                # Update document status
                await doc_service.update(doc_id, {"status": "processing"}, current_user.id)
                
                # Extract text (placeholder for MVP)
                document_text = f"""
                Sample lease agreement for {document.file_name}.
                """
                
                # Perform AI extraction
                result = await extraction_service.extract_lease_data(document_text)
                
                if not result["success"]:
                    await doc_service.update(doc_id, {"status": "failed"}, current_user.id)
                    results.append({
                        "document_id": doc_id,
                        "success": False,
                        "error": result.get("error", "Extraction failed")
                    })
                    failed += 1
                    continue
                
                extracted_data = result["data"]
                
                # Compliance check
                compliance_result = await compliance_checker.check_compliance(extracted_data)
                
                # Save extraction
                extraction = await extractions_service.create({
                    "user_id": current_user.id,
                    "document_id": doc_id,
                    "tenant_name": extracted_data.get("tenant_name"),
                    "landlord_name": extracted_data.get("landlord_name"),
                    "property_address": extracted_data.get("property_address"),
                    "monthly_rent": extracted_data.get("monthly_rent"),
                    "security_deposit": extracted_data.get("security_deposit"),
                    "lease_start_date": extracted_data.get("lease_start_date"),
                    "lease_end_date": extracted_data.get("lease_end_date"),
                    "renewal_notice_days": extracted_data.get("renewal_notice_days"),
                    "pet_policy": extracted_data.get("pet_policy"),
                    "late_fee_terms": extracted_data.get("late_fee_terms"),
                    "risk_flags": json.dumps(extracted_data.get("risk_flags", [])),
                    "compliance_data": json.dumps(compliance_result),
                    "raw_extraction": json.dumps(extracted_data),
                    "created_at": datetime.now()
                }, current_user.id)
                
                # Deduct credit (admins don't consume credits)
                if not is_admin:
                    current_credits = await credits_service.get_by_field("user_id", current_user.id)
                    if current_credits.free_credits and current_credits.free_credits > 0:
                        await credits_service.update(current_credits.id, {
                            "free_credits": current_credits.free_credits - 1,
                            "updated_at": datetime.now()
                        }, current_user.id)
                    elif current_credits.paid_credits and current_credits.paid_credits > 0:
                        await credits_service.update(current_credits.id, {
                            "paid_credits": current_credits.paid_credits - 1,
                            "updated_at": datetime.now()
                        }, current_user.id)
                
                await doc_service.update(doc_id, {"status": "completed"}, current_user.id)
                
                results.append({
                    "document_id": doc_id,
                    "success": True,
                    "extraction_id": extraction.id,
                    "file_name": document.file_name
                })
                completed += 1
                
            except Exception as e:
                logger.error(f"Batch analysis error for doc {doc_id}: {e}")
                results.append({
                    "document_id": doc_id,
                    "success": False,
                    "error": str(e)
                })
                failed += 1
        
        return BatchAnalysisResponse(
            success=completed > 0,
            results=results,
            total=len(request.document_ids),
            completed=completed,
            failed=failed
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Batch analysis error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/compliance/{extraction_id}")
async def get_compliance_report(
    extraction_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get compliance report for an extraction"""
    try:
        extractions_service = ExtractionsService(db)
        extraction = await extractions_service.get_by_id(extraction_id, current_user.id)
        
        if not extraction:
            raise HTTPException(status_code=404, detail="Extraction not found")
        
        # Re-run compliance check with current data
        compliance_checker = ComplianceChecker(db)
        extraction_data = {
            "property_address": extraction.property_address,
            "monthly_rent": extraction.monthly_rent,
            "security_deposit": extraction.security_deposit,
            "renewal_notice_days": extraction.renewal_notice_days,
            "late_fee_terms": extraction.late_fee_terms
        }
        
        compliance_result = await compliance_checker.check_compliance(extraction_data)
        
        return compliance_result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Compliance check error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/regulations/{state_code}")
async def get_state_regulations(
    state_code: str,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get all regulations for a specific state"""
    try:
        compliance_checker = ComplianceChecker(db)
        regulations = await compliance_checker.get_state_regulations(state_code.upper())
        
        if not regulations:
            raise HTTPException(
                status_code=404, 
                detail=f"No regulations found for state: {state_code}"
            )
        
        return {
            "state_code": state_code.upper(),
            "regulations": regulations
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get regulations error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/regulations")
async def get_all_states(
    db: AsyncSession = Depends(get_db),
):
    """Get list of all states with regulations"""
    try:
        from sqlalchemy import select, distinct
        from models.state_regulations import State_regulations
        
        result = await db.execute(
            select(
                State_regulations.state_code,
                State_regulations.state_name
            ).distinct()
        )
        states = result.all()
        
        return {
            "states": [
                {"code": s[0], "name": s[1]}
                for s in states
            ]
        }
        
    except Exception as e:
        logger.error(f"Get states error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/calendar/{extraction_id}")
async def get_calendar(
    extraction_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Generate ICS calendar file for lease dates"""
    try:
        extractions_service = ExtractionsService(db)
        extraction = await extractions_service.get_by_id(extraction_id, current_user.id)
        
        if not extraction:
            raise HTTPException(status_code=404, detail="Extraction not found")
        
        # Get document name
        doc_service = DocumentsService(db)
        document = await doc_service.get_by_id(extraction.document_id, current_user.id)
        doc_name = document.file_name if document else "Lease"
        
        # Generate ICS content
        extraction_data = {
            "lease_end_date": extraction.lease_end_date,
            "renewal_notice_days": extraction.renewal_notice_days,
            "property_address": extraction.property_address
        }
        
        ics_content = generate_ics_content(extraction_data, doc_name)
        
        return Response(
            content=ics_content,
            media_type="text/calendar",
            headers={
                "Content-Disposition": f"attachment; filename=lease-reminders-{extraction_id}.ics"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Calendar generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/credits")
async def get_user_credits(
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get current user's credit balance and permissions"""
    try:
        is_admin = getattr(current_user, 'role', 'user') == 'admin'
        credits_service = User_creditsService(db)
        user_credits = await credits_service.get_by_field("user_id", current_user.id)
        
        if not user_credits:
            # Create initial credits for new user
            user_credits = await credits_service.create({
                "user_id": current_user.id,
                "free_credits": 1,
                "paid_credits": 0,
                "subscription_type": "none",
                "created_at": datetime.now(),
                "updated_at": datetime.now()
            }, current_user.id)
        
        total_credits = (user_credits.free_credits or 0) + (user_credits.paid_credits or 0)
        has_subscription = False
        if user_credits.subscription_type == "monthly" and user_credits.subscription_expires_at:
            if user_credits.subscription_expires_at > datetime.now():
                has_subscription = True
        
        return {
            "free_credits": user_credits.free_credits or 0,
            "paid_credits": user_credits.paid_credits or 0,
            "total_credits": total_credits,
            "subscription_type": user_credits.subscription_type,
            "subscription_expires_at": user_credits.subscription_expires_at,
            "is_admin": is_admin,
            "can_analyze": is_admin or total_credits > 0 or has_subscription,
            "can_batch": is_admin or total_credits > 1 or has_subscription,
        }
        
    except Exception as e:
        logger.error(f"Get credits error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


class SetCreditsRequest(BaseModel):
    free_credits: Optional[int] = None
    paid_credits: Optional[int] = None
    subscription_type: Optional[str] = None


@router.post("/credits/set")
async def admin_set_credits(
    request: SetCreditsRequest,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Admin endpoint: manually set credit balance for testing"""
    is_admin = getattr(current_user, 'role', 'user') == 'admin'
    if not is_admin:
        raise HTTPException(status_code=403, detail="Admin access required")
    
    try:
        credits_service = User_creditsService(db)
        user_credits = await credits_service.get_by_field("user_id", current_user.id)
        
        if not user_credits:
            user_credits = await credits_service.create({
                "user_id": current_user.id,
                "free_credits": request.free_credits or 0,
                "paid_credits": request.paid_credits or 0,
                "subscription_type": request.subscription_type or "none",
                "created_at": datetime.now(),
                "updated_at": datetime.now()
            }, current_user.id)
        else:
            updates = {"updated_at": datetime.now()}
            if request.free_credits is not None:
                updates["free_credits"] = request.free_credits
            if request.paid_credits is not None:
                updates["paid_credits"] = request.paid_credits
            if request.subscription_type is not None:
                updates["subscription_type"] = request.subscription_type
            await credits_service.update(user_credits.id, updates, current_user.id)
        
        return {"success": True, "message": "Credits updated"}
    except Exception as e:
        logger.error(f"Set credits error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ──────────────────────────────────────────────────────────────────────────────
# Feature: PDF Source Tracing (click-to-source)
# ──────────────────────────────────────────────────────────────────────────────


@router.get("/source-map/{extraction_id}")
async def get_source_map(
    extraction_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get source location mapping for an extraction (field → PDF location)."""
    try:
        extractions_service = ExtractionsService(db)
        extraction = await extractions_service.get_by_id(extraction_id, current_user.id)

        if not extraction:
            raise HTTPException(status_code=404, detail="Extraction not found")

        source_map = {}
        if extraction.source_map:
            try:
                source_map = json.loads(extraction.source_map)
            except (json.JSONDecodeError, TypeError):
                pass

        pages_meta = []
        if extraction.pages_meta:
            try:
                pages_meta = json.loads(extraction.pages_meta)
            except (json.JSONDecodeError, TypeError):
                pass

        return {
            "extraction_id": extraction_id,
            "source_map": source_map,
            "pages_meta": pages_meta,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Source map error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/pdf-page/{document_id}/{page_num}")
async def get_pdf_page_image(
    document_id: int,
    page_num: int,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Render a specific document page as a PNG image for the viewer.
    Supports both PDF and Word (.docx) files."""
    try:
        doc_service = DocumentsService(db)
        document = await doc_service.get_by_id(document_id, current_user.id)

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # Download file from storage
        storage = StorageService()
        download = await storage.create_download_url(
            FileUpDownRequest(bucket_name="lease-documents", object_key=document.file_key)
        )
        async with httpx_client.AsyncClient(timeout=120.0) as http_client:
            file_response = await http_client.get(download.download_url)
            file_bytes = file_response.content

        file_name = document.file_name.lower()

        if file_name.endswith(('.docx', '.doc')):
            # Convert Word to PDF in memory, then render page
            img_bytes = _render_docx_page(file_bytes, page_num)
        else:
            # Render PDF page
            pdf_extractor = PDFExtractor()
            img_bytes = pdf_extractor.get_page_image(file_bytes, page_num, dpi=150)

        return Response(
            content=img_bytes,
            media_type="image/png",
            headers={"Cache-Control": "public, max-age=3600"},
        )

    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Document page render error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _render_docx_page(file_bytes: bytes, page_num: int) -> bytes:
    """Render a Word document page as PNG by converting to PDF first via reportlab,
    or by rendering text content as an image."""
    import io
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ValueError("python-docx not installed")

    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    # Paginate: ~45 lines per page (standard letter)
    lines_per_page = 45
    pages = []
    current_page_lines = []
    for para in paragraphs:
        # Wrap long paragraphs
        words = para.split()
        line = ""
        for word in words:
            test = f"{line} {word}".strip()
            if len(test) > 85:  # ~85 chars per line at 11pt
                current_page_lines.append(line)
                line = word
                if len(current_page_lines) >= lines_per_page:
                    pages.append(current_page_lines)
                    current_page_lines = []
            else:
                line = test
        if line:
            current_page_lines.append(line)
            if len(current_page_lines) >= lines_per_page:
                pages.append(current_page_lines)
                current_page_lines = []
        current_page_lines.append("")  # blank line between paragraphs

    if current_page_lines:
        pages.append(current_page_lines)

    if not pages:
        pages = [["(Empty document)"]]

    if page_num >= len(pages):
        raise ValueError(f"Page {page_num} does not exist (total pages: {len(pages)})")

    # Render page as image using reportlab
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas as rl_canvas
    import fitz

    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 72  # 1 inch top margin

    for line in pages[page_num]:
        if y < 72:
            break
        c.setFont("Helvetica", 11)
        c.drawString(72, y, line)
        y -= 16

    c.save()
    buf.seek(0)

    # Convert single-page PDF to PNG
    pdf_doc = fitz.open(stream=buf.read(), filetype="pdf")
    page = pdf_doc[0]
    zoom = 150 / 72
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat)
    img_bytes = pix.tobytes("png")
    pdf_doc.close()
    return img_bytes


@router.get("/doc-page-count/{document_id}")
async def get_document_page_count(
    document_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get total page count for a document (PDF or Word)."""
    try:
        doc_service = DocumentsService(db)
        document = await doc_service.get_by_id(document_id, current_user.id)

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        storage = StorageService()
        download = await storage.create_download_url(
            FileUpDownRequest(bucket_name="lease-documents", object_key=document.file_key)
        )
        async with httpx_client.AsyncClient(timeout=120.0) as http_client:
            file_response = await http_client.get(download.download_url)
            file_bytes = file_response.content

        file_name = document.file_name.lower()

        if file_name.endswith(('.docx', '.doc')):
            import io
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Estimate pages
            total_lines = 0
            for para in paragraphs:
                total_lines += max(1, len(para) // 85 + 1) + 1
            page_count = max(1, (total_lines + 44) // 45)
        else:
            import fitz
            pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
            page_count = len(pdf_doc)
            pdf_doc.close()

        return {"page_count": page_count, "file_type": "docx" if file_name.endswith(('.docx', '.doc')) else "pdf"}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Page count error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/pdf-url/{document_id}")
async def get_pdf_download_url(
    document_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get a temporary download URL for the original PDF (for pdfjs viewer)."""
    try:
        doc_service = DocumentsService(db)
        document = await doc_service.get_by_id(document_id, current_user.id)

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        storage = StorageService()
        download = await storage.create_download_url(
            FileUpDownRequest(bucket_name="lease-documents", object_key=document.file_key)
        )

        return {"url": download.download_url}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PDF URL error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ──────────────────────────────────────────────────────────────────────────────
# Feature: Google Calendar Integration
# ──────────────────────────────────────────────────────────────────────────────


@router.get("/google-calendar/{extraction_id}")
async def get_google_calendar_links(
    extraction_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Generate Google Calendar event links for key lease dates."""
    try:
        extractions_service = ExtractionsService(db)
        extraction = await extractions_service.get_by_id(extraction_id, current_user.id)

        if not extraction:
            raise HTTPException(status_code=404, detail="Extraction not found")

        doc_service = DocumentsService(db)
        document = await doc_service.get_by_id(extraction.document_id, current_user.id)
        doc_name = document.file_name if document else "Lease"

        events: List[Dict[str, Any]] = []

        property_addr = extraction.property_address or "N/A"

        # Lease end date event
        if extraction.lease_end_date:
            date_str = extraction.lease_end_date.replace("-", "")
            events.append({
                "title": f"Lease Expires – {doc_name}",
                "date": extraction.lease_end_date,
                "google_url": _build_gcal_url(
                    title=f"Lease Expires – {doc_name}",
                    date=date_str,
                    details=f"Your lease agreement expires. Property: {property_addr}",
                    location=property_addr,
                ),
                "type": "lease_end",
            })

            # Renewal notice deadline
            notice_days = extraction.renewal_notice_days or 60
            try:
                from datetime import timedelta
                end_date = datetime.strptime(extraction.lease_end_date, "%Y-%m-%d")
                notice_date = end_date - timedelta(days=notice_days)
                notice_str = notice_date.strftime("%Y%m%d")
                events.append({
                    "title": f"Renewal Notice Deadline – {doc_name}",
                    "date": notice_date.strftime("%Y-%m-%d"),
                    "google_url": _build_gcal_url(
                        title=f"Renewal Notice Deadline – {doc_name}",
                        date=notice_str,
                        details=f"Deadline to provide renewal notice ({notice_days} days before lease end). Property: {property_addr}",
                        location=property_addr,
                    ),
                    "type": "renewal_notice",
                })
            except ValueError:
                pass

        # Lease start date event (for reference)
        if extraction.lease_start_date:
            date_str = extraction.lease_start_date.replace("-", "")
            events.append({
                "title": f"Lease Starts – {doc_name}",
                "date": extraction.lease_start_date,
                "google_url": _build_gcal_url(
                    title=f"Lease Starts – {doc_name}",
                    date=date_str,
                    details=f"Lease agreement start date. Property: {property_addr}",
                    location=property_addr,
                ),
                "type": "lease_start",
            })

        return {"events": events}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Google Calendar link error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _build_gcal_url(title: str, date: str, details: str, location: str = "") -> str:
    """Build a Google Calendar event creation URL."""
    params = {
        "action": "TEMPLATE",
        "text": title,
        "dates": f"{date}/{date}",
        "details": details,
        "location": location,
    }
    return f"https://calendar.google.com/calendar/render?{urlencode(params)}"


# ──────────────────────────────────────────────────────────────────────────────
# Feature: Amendment Memo Generation
# ──────────────────────────────────────────────────────────────────────────────


@router.post("/amendment-memo/{extraction_id}")
async def generate_amendment_memo(
    extraction_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Generate a professional amendment suggestions memo for a lease."""
    try:
        extractions_service = ExtractionsService(db)
        extraction = await extractions_service.get_by_id(extraction_id, current_user.id)

        if not extraction:
            raise HTTPException(status_code=404, detail="Extraction not found")

        # Build extraction data
        extraction_data = {
            "tenant_name": extraction.tenant_name,
            "landlord_name": extraction.landlord_name,
            "property_address": extraction.property_address,
            "monthly_rent": extraction.monthly_rent,
            "security_deposit": extraction.security_deposit,
            "lease_start_date": extraction.lease_start_date,
            "lease_end_date": extraction.lease_end_date,
            "pet_policy": extraction.pet_policy,
            "late_fee_terms": extraction.late_fee_terms,
            "renewal_notice_days": extraction.renewal_notice_days,
        }

        # Parse risk flags safely (no eval)
        risk_flags = []
        if extraction.risk_flags:
            risk_flags = safe_parse_json_or_literal(extraction.risk_flags)
            if not isinstance(risk_flags, list):
                risk_flags = []

        # Parse compliance data safely (no eval)
        compliance_data = None
        if extraction.compliance_data:
            compliance_data = safe_parse_json_or_literal(extraction.compliance_data)

        memo_service = AmendmentMemoService()
        result = await memo_service.generate_memo(extraction_data, compliance_data, risk_flags)

        return result

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Amendment memo error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ──────────────────────────────────────────────────────────────────────────────
# Feature: Lease Comparison
# ──────────────────────────────────────────────────────────────────────────────


class CompareRequest(BaseModel):
    extraction_ids: List[int]


@router.post("/compare")
async def compare_leases(
    request: CompareRequest,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Compare 2-3 lease extractions side by side."""
    try:
        if len(request.extraction_ids) < 2:
            raise HTTPException(status_code=400, detail="Need at least 2 extractions to compare")
        if len(request.extraction_ids) > 3:
            raise HTTPException(status_code=400, detail="Can compare up to 3 leases at a time")

        extractions_service = ExtractionsService(db)
        extractions_data = []
        labels = []

        for ext_id in request.extraction_ids:
            extraction = await extractions_service.get_by_id(ext_id, current_user.id)
            if not extraction:
                raise HTTPException(status_code=404, detail=f"Extraction {ext_id} not found")

            ext_data = {
                "tenant_name": extraction.tenant_name,
                "landlord_name": extraction.landlord_name,
                "property_address": extraction.property_address,
                "monthly_rent": extraction.monthly_rent,
                "security_deposit": extraction.security_deposit,
                "lease_start_date": extraction.lease_start_date,
                "lease_end_date": extraction.lease_end_date,
                "renewal_notice_days": extraction.renewal_notice_days,
                "pet_policy": extraction.pet_policy,
                "late_fee_terms": extraction.late_fee_terms,
                "risk_flags": extraction.risk_flags,
            }
            extractions_data.append(ext_data)
            labels.append(extraction.property_address or f"Lease #{ext_id}")

        comparison_service = LeaseComparisonService()
        result = await comparison_service.compare_leases(extractions_data, labels)

        return result

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Lease comparison error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ──────────────────────────────────────────────────────────────────────────────
# Feature: Rent Benchmarking
# ──────────────────────────────────────────────────────────────────────────────


@router.get("/benchmark/{extraction_id}")
async def get_rent_benchmark(
    extraction_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get market rent benchmark data for a lease."""
    try:
        extractions_service = ExtractionsService(db)
        extraction = await extractions_service.get_by_id(extraction_id, current_user.id)

        if not extraction:
            raise HTTPException(status_code=404, detail="Extraction not found")

        benchmark_service = RentBenchmarkService()
        result = benchmark_service.get_benchmark(
            property_address=extraction.property_address or "",
            monthly_rent=extraction.monthly_rent,
        )

        return {"success": True, "benchmark": result}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Benchmark error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ──────────────────────────────────────────────────────────────────────────────
# Feature: Email Lead Capture (Template Downloads)
# ──────────────────────────────────────────────────────────────────────────────


class LeadCaptureRequest(BaseModel):
    email: str
    name: Optional[str] = None
    template_type: str = "general"


@router.post("/lead/capture")
async def capture_lead(request: LeadCaptureRequest):
    """Capture email for template downloads (no auth required)."""
    try:
        # In production, save to DB or send to email marketing service
        logger.info(f"Lead captured: {request.email} for template: {request.template_type}")
        return {
            "success": True,
            "message": "Thank you! Check your email for the download link.",
            "download_url": f"/api/v1/lease/template/{request.template_type}",
        }
    except Exception as e:
        logger.error(f"Lead capture error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/template/{template_type}")
async def get_template(template_type: str):
    """Return template content (placeholder - in production serve actual files)."""
    templates = {
        "checklist": {
            "title": "Landlord Lease Review Checklist",
            "type": "checklist",
            "items": [
                "Verify tenant full legal name matches ID",
                "Confirm property address is complete and correct",
                "Check security deposit amount complies with state law",
                "Ensure lease start and end dates are clearly specified",
                "Review late fee terms for reasonableness and legality",
                "Confirm pet policy is explicitly stated",
                "Check for lead-based paint disclosure (pre-1978 buildings)",
                "Verify maintenance responsibilities are clearly defined",
                "Review subletting/assignment clauses",
                "Confirm notice periods for entry, termination, and renewal",
                "Check for dispute resolution procedures",
                "Verify insurance requirements",
                "Review early termination clauses and penalties",
                "Confirm utility responsibility assignments",
                "Check parking and storage arrangements",
            ],
        },
        "california": {
            "title": "California Residential Lease Key Requirements",
            "type": "guide",
            "sections": [
                {"title": "Security Deposit", "content": "Maximum 2x monthly rent (unfurnished) or 3x (furnished). AB 12 reduces this to 1x starting 2025."},
                {"title": "Late Fees", "content": "Must be reasonable. Courts generally accept 5-6% of monthly rent."},
                {"title": "Rent Increase Notice", "content": "30 days for increases <10%, 90 days for increases ≥10%."},
                {"title": "Entry Notice", "content": "24 hours written notice required except emergencies."},
                {"title": "Rent Control", "content": "AB 1482 caps annual increases at 5% + CPI (max 10%) for properties 15+ years old."},
            ],
        },
        "general": {
            "title": "Standard Residential Lease Template Guide",
            "type": "guide",
            "sections": [
                {"title": "Essential Parties", "content": "Full legal names of all tenants and landlord/property management company."},
                {"title": "Property Description", "content": "Complete address including unit number, parking spaces, storage areas."},
                {"title": "Term", "content": "Clear start and end dates, renewal terms, month-to-month conversion."},
                {"title": "Rent", "content": "Amount, due date, acceptable payment methods, grace period."},
                {"title": "Security Deposit", "content": "Amount, conditions for return, timeline, itemization requirements."},
                {"title": "Maintenance", "content": "Landlord vs tenant responsibilities, repair request procedures, emergency contacts."},
            ],
        },
    }
    return templates.get(template_type, templates["general"])


# ──────────────────────────────────────────────────────────────────────────────
# Feature: Portfolio Summary
# ──────────────────────────────────────────────────────────────────────────────


@router.get("/portfolio")
async def get_portfolio_summary(
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get portfolio-level summary of all user's leases."""
    try:
        extractions_service = ExtractionsService(db)
        doc_service = DocumentsService(db)

        # Get all user documents
        from sqlalchemy import select
        from models.documents import Documents
        from models.extractions import Extractions

        docs_result = await db.execute(
            select(Documents).where(Documents.created_by == current_user.id).order_by(Documents.created_at.desc())
        )
        documents = docs_result.scalars().all()

        ext_result = await db.execute(
            select(Extractions).where(Extractions.user_id == current_user.id).order_by(Extractions.created_at.desc())
        )
        extractions = ext_result.scalars().all()

        # Build portfolio data
        total_monthly_rent = 0
        total_deposits = 0
        active_leases = 0
        expiring_soon = []  # within 90 days
        upcoming_dates = []

        from datetime import datetime, timedelta
        now = datetime.now()
        ninety_days = now + timedelta(days=90)

        lease_items = []
        for ext in extractions:
            item = {
                "extraction_id": ext.id,
                "document_id": ext.document_id,
                "property_address": ext.property_address,
                "tenant_name": ext.tenant_name,
                "monthly_rent": ext.monthly_rent,
                "security_deposit": ext.security_deposit,
                "lease_start_date": ext.lease_start_date,
                "lease_end_date": ext.lease_end_date,
                "status": "active",
            }

            if ext.monthly_rent:
                total_monthly_rent += ext.monthly_rent
            if ext.security_deposit:
                total_deposits += ext.security_deposit

            # Check lease status
            if ext.lease_end_date:
                try:
                    end_date = datetime.strptime(ext.lease_end_date, "%Y-%m-%d")
                    if end_date < now:
                        item["status"] = "expired"
                    elif end_date <= ninety_days:
                        item["status"] = "expiring_soon"
                        expiring_soon.append(item)
                    else:
                        item["status"] = "active"
                        active_leases += 1

                    upcoming_dates.append({
                        "date": ext.lease_end_date,
                        "type": "lease_end",
                        "label": f"Lease ends – {ext.property_address or 'Unknown'}",
                        "extraction_id": ext.id,
                    })
                except ValueError:
                    item["status"] = "unknown"
            else:
                active_leases += 1

            lease_items.append(item)

        # Sort upcoming dates
        upcoming_dates.sort(key=lambda x: x["date"])

        return {
            "summary": {
                "total_properties": len(lease_items),
                "active_leases": active_leases,
                "expiring_soon": len(expiring_soon),
                "total_monthly_rent": round(total_monthly_rent, 2),
                "total_annual_rent": round(total_monthly_rent * 12, 2),
                "total_deposits_held": round(total_deposits, 2),
            },
            "leases": lease_items,
            "upcoming_dates": upcoming_dates[:10],
            "expiring_soon": expiring_soon,
        }

    except Exception as e:
        logger.error(f"Portfolio summary error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ──────────────────────────────────────────────────────────────────────────────
# Feature: Executive Summary (AI)
# ──────────────────────────────────────────────────────────────────────────────


@router.get("/summary/{extraction_id}")
async def get_executive_summary(
    extraction_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Generate an AI executive summary for a lease report."""
    try:
        extractions_service = ExtractionsService(db)
        extraction = await extractions_service.get_by_id(extraction_id, current_user.id)
        if not extraction:
            raise HTTPException(status_code=404, detail="Extraction not found")

        extraction_data = {
            "tenant_name": extraction.tenant_name,
            "landlord_name": extraction.landlord_name,
            "property_address": extraction.property_address,
            "monthly_rent": extraction.monthly_rent,
            "security_deposit": extraction.security_deposit,
            "lease_start_date": extraction.lease_start_date,
            "lease_end_date": extraction.lease_end_date,
            "pet_policy": extraction.pet_policy,
            "late_fee_terms": extraction.late_fee_terms,
            "renewal_notice_days": extraction.renewal_notice_days,
        }

        # Parse risk flags safely (no eval)
        risk_flags = []
        if extraction.risk_flags:
            risk_flags = safe_parse_json_or_literal(extraction.risk_flags)
            if not isinstance(risk_flags, list):
                risk_flags = []

        # Parse compliance data safely (no eval)
        compliance_data = None
        if extraction.compliance_data:
            compliance_data = safe_parse_json_or_literal(extraction.compliance_data)

        summary_service = ExecutiveSummaryService()
        result = await summary_service.generate_summary(extraction_data, compliance_data, risk_flags)
        return result

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Executive summary error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ──────────────────────────────────────────────────────────────────────────────
# Feature: PDF Report Export
# ──────────────────────────────────────────────────────────────────────────────


@router.get("/export-pdf/{extraction_id}")
async def export_pdf_report(
    extraction_id: int,
    current_user: UserResponse = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Generate and return a professional PDF report."""
    try:
        extractions_service = ExtractionsService(db)
        extraction = await extractions_service.get_by_id(extraction_id, current_user.id)
        if not extraction:
            raise HTTPException(status_code=404, detail="Extraction not found")

        extraction_data = {
            "tenant_name": extraction.tenant_name,
            "landlord_name": extraction.landlord_name,
            "property_address": extraction.property_address,
            "monthly_rent": extraction.monthly_rent,
            "security_deposit": extraction.security_deposit,
            "lease_start_date": extraction.lease_start_date,
            "lease_end_date": extraction.lease_end_date,
            "pet_policy": extraction.pet_policy,
            "late_fee_terms": extraction.late_fee_terms,
            "renewal_notice_days": extraction.renewal_notice_days,
        }

        # Parse risk flags safely (no eval)
        risk_flags = []
        if extraction.risk_flags:
            risk_flags = safe_parse_json_or_literal(extraction.risk_flags)
            if not isinstance(risk_flags, list):
                risk_flags = []

        # Parse compliance data safely (no eval)
        compliance_data = None
        if extraction.compliance_data:
            compliance_data = safe_parse_json_or_literal(extraction.compliance_data)

        # Get benchmark data
        from services.rent_benchmark import RentBenchmarkService
        benchmark_service = RentBenchmarkService()
        benchmark_data = benchmark_service.get_benchmark(
            property_address=extraction.property_address or "",
            monthly_rent=extraction.monthly_rent,
        )

        # Get executive summary
        try:
            summary_service = ExecutiveSummaryService()
            executive_summary = await summary_service.generate_summary(extraction_data, compliance_data, risk_flags)
        except Exception:
            executive_summary = None

        # Generate PDF
        generator = PDFReportGenerator()
        pdf_bytes = generator.generate(
            extraction=extraction_data,
            risk_flags=risk_flags,
            compliance_data=compliance_data,
            benchmark_data=benchmark_data,
            executive_summary=executive_summary,
        )

        if not pdf_bytes:
            raise HTTPException(status_code=500, detail="PDF generation failed. reportlab may not be installed.")

        filename = f"LeaseLens-Report-{extraction_id}.pdf"
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PDF export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))